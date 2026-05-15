#!/usr/bin/env python3
"""
Bulk-translate large Markdown or DOCX files with fewer round-trips than per-fence loops.

Engines (no paid API keys required):
  - argos  — Argos Translate (offline, local). Default. Auto-installs missing pairs on first use.
  - google — deep_translator / Google (free tier). Use --batch-merge-chars to merge many segments
             into fewer HTTP calls (much faster than one request per ```text block).

Language codes (normalized):
  en, vi, zh (Simplified Chinese), zt (Traditional Chinese in Argos).

Pairs without a direct Argos model are chained via English (e.g. vi → zh: vi → en → zh).

Examples:
  python scripts/bulk_translate_file.py -i FUNC_IMPROVE/foo.md -o FUNC_IMPROVE/foo_en.md -s vi -t en
  python scripts/bulk_translate_file.py -i report.md -s vi -t zh --engine argos
  python scripts/bulk_translate_file.py -i big.md -s vi -t en --engine google --batch-merge-chars 4500
  python scripts/bulk_translate_file.py -i notes.docx -s en -t vi -o notes_vi.docx
"""

from __future__ import annotations

import argparse
import re
import sys
import time
from pathlib import Path

# ---------------------------------------------------------------------------
# Language normalization
# ---------------------------------------------------------------------------

LANG_ALIASES: dict[str, str] = {
    "zh-cn": "zh",
    "zh_cn": "zh",
    "zh-hans": "zh",
    "cmn": "zh",
    "zh-tw": "zt",
    "zh_tw": "zt",
    "zh-hant": "zt",
    "yue": "zt",
}


def normalize_lang(code: str) -> str:
    c = code.strip().lower().replace("_", "-")
    return LANG_ALIASES.get(c, c.split("-")[0] if c.startswith("zh-") else c)


# ASCII delimiters for Google batch merge (easier to regex-split than Unicode).
def _merge_marker(idx: int) -> str:
    return f"\n<<<__BTX_{idx:06d}__>>>\n"


def _split_google_merged(translated: str, batch_ids: list[int]) -> dict[int, str]:
    """Split translated blob back into per-index strings using <<<__BTX_NNNNNN__>>> markers."""
    pattern = re.compile(r"\n<<<__BTX_(\d{6})__>>>\n")
    parts = pattern.split(translated)
    out: dict[int, str] = {}
    if not batch_ids:
        return out
    if len(parts) == 1:
        out[batch_ids[0]] = parts[0].strip()
        return out
    if parts[0].strip():
        out[batch_ids[0]] = parts[0].strip()
    for i in range(1, len(parts), 2):
        try:
            idx = int(parts[i])
        except (ValueError, IndexError):
            continue
        body = parts[i + 1] if i + 1 < len(parts) else ""
        out[idx] = body.strip()
    return out


# ---------------------------------------------------------------------------
# Argos
# ---------------------------------------------------------------------------

def _argos_install_pair(from_code: str, to_code: str) -> None:
    import argostranslate.package as pkg

    pkg.update_package_index()
    ap = next((p for p in pkg.get_available_packages() if p.from_code == from_code and p.to_code == to_code), None)
    if ap is None:
        raise RuntimeError(f"No Argos package published for {from_code} → {to_code}.")
    print(f"  [argos] Installing package {from_code} → {to_code} …", flush=True)
    path = ap.download()
    pkg.install_from_path(path)


def _argos_has_pair(from_code: str, to_code: str) -> bool:
    from argostranslate.translate import get_installed_languages

    for lang in get_installed_languages():
        if lang.code != from_code:
            continue
        for t in lang.translations_from:
            if t.to_lang.code == to_code:
                return True
    return False


def ensure_argos_pair(from_code: str, to_code: str, *, allow_install: bool) -> None:
    if _argos_has_pair(from_code, to_code):
        return
    if not allow_install:
        raise RuntimeError(f"Argos pair missing {from_code}→{to_code}. Run with install enabled or install manually.")
    _argos_install_pair(from_code, to_code)


def translate_argos_piece(text: str, from_code: str, to_code: str) -> str:
    from argostranslate.translate import translate

    return translate(text, from_code, to_code)


def translate_argos(text: str, from_code: str, to_code: str, *, allow_install: bool) -> str:
    if from_code == to_code or not text.strip():
        return text

    def one_step(src: str, tgt: str, body: str) -> str:
        ensure_argos_pair(src, tgt, allow_install=allow_install)
        max_chunk = 3500
        if len(body) <= max_chunk:
            return translate_argos_piece(body, src, tgt)
        chunks: list[str] = []
        start = 0
        while start < len(body):
            end = min(start + max_chunk, len(body))
            if end < len(body):
                cut = body.rfind("\n\n", start, end)
                if cut > start + 500:
                    end = cut
            piece = body[start:end]
            chunks.append(translate_argos_piece(piece, src, tgt))
            start = end
        return "\n\n".join(chunks)

    if _argos_has_pair(from_code, to_code) or allow_install:
        try:
            ensure_argos_pair(from_code, to_code, allow_install=allow_install)
            return one_step(from_code, to_code, text)
        except RuntimeError:
            pass

    if from_code != "en" and to_code != "en":
        ensure_argos_pair(from_code, "en", allow_install=allow_install)
        ensure_argos_pair("en", to_code, allow_install=allow_install)
        mid = one_step(from_code, "en", text)
        return one_step("en", to_code, mid)

    raise RuntimeError(f"Cannot resolve Argos path {from_code} → {to_code}")


# ---------------------------------------------------------------------------
# Google (deep_translator) with optional batch merge
# ---------------------------------------------------------------------------

def translate_google(text: str, from_code: str, to_code: str) -> str:
    from deep_translator import GoogleTranslator

    src = "auto" if from_code == "auto" else from_code
    translator = GoogleTranslator(source=src, target=to_code)
    out: list[str] = []
    for i in range(0, len(text), 4500):
        chunk = text[i : i + 4500]
        try:
            out.append(translator.translate(chunk))
        except Exception:
            out.append(chunk)
        time.sleep(0.05)
    return "".join(out)


def translate_google_batched(segments: list[str], from_code: str, to_code: str, merge_chars: int) -> list[str]:
    if not segments:
        return []
    from deep_translator import GoogleTranslator

    src = "auto" if from_code == "auto" else from_code
    translator = GoogleTranslator(source=src, target=to_code)
    n = len(segments)
    results: list[str] = [""] * n

    def translate_single(text: str) -> str:
        return translate_google(text, from_code, to_code)

    i = 0
    while i < n:
        batch_ids: list[int] = []
        size = 0
        j = i
        while j < n:
            seg = segments[j]
            overhead = len(_merge_marker(j)) if batch_ids else 0
            if batch_ids and size + overhead + len(seg) > merge_chars:
                break
            if not batch_ids:
                size = len(seg)
            else:
                size += overhead + len(seg)
            batch_ids.append(j)
            j += 1
        if not batch_ids:
            batch_ids = [i]
            j = i + 1

        if len(batch_ids) == 1:
            idx = batch_ids[0]
            body = segments[idx]
            if len(body) > merge_chars:
                results[idx] = translate_single(body)
            else:
                try:
                    results[idx] = translator.translate(body).strip()
                except Exception:
                    results[idx] = body
                time.sleep(0.05)
        else:
            merged = segments[batch_ids[0]]
            for idx in batch_ids[1:]:
                merged += _merge_marker(idx) + segments[idx]
            try:
                translated = translator.translate(merged[:50000])
            except Exception:
                translated = merged
            time.sleep(0.06)
            mapping = _split_google_merged(translated, batch_ids)
            for idx in batch_ids:
                if idx in mapping and mapping[idx]:
                    results[idx] = mapping[idx]
                else:
                    try:
                        results[idx] = translator.translate(segments[idx]).strip()
                    except Exception:
                        results[idx] = segments[idx]
                    time.sleep(0.04)

        i = batch_ids[-1] + 1

    return results


# ---------------------------------------------------------------------------
# Markdown modes
# ---------------------------------------------------------------------------

MD_TEXT_FENCE = re.compile(r"```text\r?\n(.*?)```", re.DOTALL)


def translate_markdown_fenced(
    raw: str,
    *,
    engine: str,
    from_lang: str,
    to_lang: str,
    batch_merge_chars: int,
    allow_argos_install: bool,
) -> str:
    matches = list(MD_TEXT_FENCE.finditer(raw))
    originals = [m.group(1) for m in matches]
    if not originals:
        return raw

    if engine == "argos":
        results = [translate_argos(t, from_lang, to_lang, allow_install=allow_argos_install) for t in originals]
    else:
        if batch_merge_chars > 800 and len(originals) > 3:
            results = translate_google_batched(originals, from_lang, to_lang, batch_merge_chars)
        else:
            results = [translate_google(t, from_lang, to_lang) for t in originals]

    out: list[str] = []
    pos = 0
    for i, m in enumerate(matches):
        out.append(raw[pos : m.start()])
        nl = "\n" if "\r\n" not in m.group(0) else "\r\n"
        out.append(f"```text{nl}")
        out.append(results[i])
        out.append(f"{nl}```")
        pos = m.end()
    out.append(raw[pos:])
    return "".join(out)


def translate_markdown_whole(
    raw: str,
    *,
    engine: str,
    from_lang: str,
    to_lang: str,
    batch_merge_chars: int,
    allow_argos_install: bool,
) -> str:
    if engine == "argos":
        return translate_argos(raw, from_lang, to_lang, allow_install=allow_argos_install)
    if len(raw) <= batch_merge_chars:
        return translate_google(raw, from_lang, to_lang)
    parts = translate_google_batched([raw], from_lang, to_lang, batch_merge_chars)
    return parts[0] if parts else raw


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def translate_docx(
    path_in: Path,
    path_out: Path,
    *,
    engine: str,
    from_lang: str,
    to_lang: str,
    batch_merge_chars: int,
    allow_argos_install: bool,
) -> None:
    from docx import Document

    doc = Document(str(path_in))
    paras = [p.text for p in doc.paragraphs]
    if engine == "argos":
        translated = [translate_argos(t, from_lang, to_lang, allow_install=allow_argos_install) for t in paras]
    else:
        translated = translate_google_batched(paras, from_lang, to_lang, min(batch_merge_chars, 4000))

    if len(paras) != len(translated):
        raise RuntimeError("DOCX paragraph count mismatch after translation.")
    for p, new in zip(doc.paragraphs, translated, strict=True):
        p.text = new

    doc.save(str(path_out))


# ---------------------------------------------------------------------------
# Self-test
# ---------------------------------------------------------------------------

def run_selftest() -> int:
    vi = "Xin chào. Đây là câu thử."
    en = translate_argos(vi, "vi", "en", allow_install=True)
    low = en.lower()
    if "xin" in low and "chào" in en:
        print("Self-test Argos vi→en looks untranslated:", repr(en), file=sys.stderr)
        return 1
    md = "```text\n" + vi + "\n```\n\nMore **bold**.\n"
    out = translate_markdown_fenced(
        md, engine="argos", from_lang="vi", to_lang="en", batch_merge_chars=4000, allow_argos_install=True
    )
    if "```text" not in out:
        print("Self-test fence missing", file=sys.stderr)
        return 1
    print("Self-test OK (Argos vi→en + markdown fence). Sample:", repr(en)[:80])
    return 0


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("-i", "--input", type=Path, default=None, help="Input file (.md, .markdown, .docx)")
    p.add_argument("-o", "--output", type=Path, default=None, help="Output path (default: sibling with _<target>.<ext>)")
    p.add_argument("-s", "--source", default=None, help="Source language: en | vi | zh | zt | auto (google only)")
    p.add_argument("-t", "--target", default=None, help="Target language: en | vi | zh | zt")
    p.add_argument(
        "--engine",
        choices=("argos", "google"),
        default="argos",
        help="Translation engine (default: argos, offline)",
    )
    p.add_argument(
        "--mode",
        choices=("md-text-fences", "md-whole"),
        default="md-text-fences",
        help="For Markdown: only ```text fences, or translate entire file (default: md-text-fences)",
    )
    p.add_argument(
        "--batch-merge-chars",
        type=int,
        default=4200,
        help="For --engine google: merge this many chars of ```text bodies per request (default: 4200)",
    )
    p.add_argument(
        "--no-argos-install",
        action="store_true",
        help="Do not auto-download Argos language packages (fail if missing)",
    )
    p.add_argument("--selftest", action="store_true", help="Run quick offline check and exit")
    args = p.parse_args()

    if args.selftest:
        return run_selftest()

    if args.input is None:
        p.error("--input is required unless --selftest is used.")
    if args.source is None or args.target is None:
        p.error("--source and --target are required unless --selftest is used.")

    path_in: Path = args.input.resolve()
    if not path_in.is_file():
        print(f"Input not found: {path_in}", file=sys.stderr)
        return 1

    ext = path_in.suffix.lower()
    if ext not in {".md", ".markdown", ".docx"}:
        print("Supported inputs: .md, .markdown, .docx (legacy .doc: save as .docx first).", file=sys.stderr)
        return 1

    src = normalize_lang(args.source)
    tgt = normalize_lang(args.target)
    if src == tgt:
        print("Source and target are the same; nothing to do.", file=sys.stderr)
        return 1
    if args.engine == "argos" and src == "auto":
        print("--source auto is only supported with --engine google.", file=sys.stderr)
        return 1

    path_out = args.output
    if path_out is None:
        path_out = path_in.with_name(f"{path_in.stem}_{tgt}{path_in.suffix}")

    allow_install = not args.no_argos_install

    print(f"Input:  {path_in}", flush=True)
    print(f"Output: {path_out}", flush=True)
    print(f"Engine: {args.engine}  {src} → {tgt}  mode={args.mode}", flush=True)

    if ext == ".docx":
        path_out.parent.mkdir(parents=True, exist_ok=True)
        translate_docx(
            path_in,
            path_out,
            engine=args.engine,
            from_lang=src,
            to_lang=tgt,
            batch_merge_chars=args.batch_merge_chars,
            allow_argos_install=allow_install,
        )
    else:
        raw = path_in.read_text(encoding="utf-8")
        if args.mode == "md-text-fences":
            new_raw = translate_markdown_fenced(
                raw,
                engine=args.engine,
                from_lang=src,
                to_lang=tgt,
                batch_merge_chars=args.batch_merge_chars,
                allow_argos_install=allow_install,
            )
        else:
            new_raw = translate_markdown_whole(
                raw,
                engine=args.engine,
                from_lang=src,
                to_lang=tgt,
                batch_merge_chars=args.batch_merge_chars,
                allow_argos_install=allow_install,
            )
        path_out.parent.mkdir(parents=True, exist_ok=True)
        path_out.write_text(new_raw, encoding="utf-8")

    print("Done.", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
