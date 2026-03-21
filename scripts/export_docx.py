#!/usr/bin/env python3
"""Convert Markdown (headings, pipe tables, ** / `) to .docx via python-docx."""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_runs_with_markup(paragraph, text: str) -> None:
    if not text:
        return
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def is_table_separator(line: str) -> bool:
    line = line.strip()
    if not line.startswith("|"):
        return False
    inner = line.strip("|").strip()
    return bool(re.match(r"^[\s|\-:]+$", inner))


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        line = line.rstrip()
        if not line.strip().startswith("|"):
            continue
        if is_table_separator(line):
            continue
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    return rows


def md_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            rows = parse_table(tbl_lines)
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = t.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs_with_markup(p, cell_text)
                        if ri == 0:
                            for r in p.runs:
                                r.bold = True
                doc.add_paragraph()
            continue

        para_lines: list[str] = []
        while i < n:
            s = lines[i].strip()
            if not s:
                break
            if s == "---":
                break
            if s.startswith("#"):
                break
            if s.startswith("|"):
                break
            para_lines.append(s)
            i += 1
        text = " ".join(para_lines)
        if text:
            p = doc.add_paragraph()
            add_runs_with_markup(p, text)
            p.paragraph_format.space_after = Pt(6)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


def main() -> None:
    ap = argparse.ArgumentParser(description="Export Markdown to Word .docx")
    ap.add_argument("input_md", type=Path, nargs="?", default=None)
    ap.add_argument("-o", "--output", type=Path, default=None)
    args = ap.parse_args()
    root = Path(__file__).resolve().parents[1]
    md = (args.input_md or (root / "docs" / "CRAVEVA_PARTNER_TECH_SPEC.md")).resolve()
    if not md.is_file():
        print(f"Not found: {md}", file=sys.stderr)
        sys.exit(1)
    out = args.output or md.with_suffix(".docx")
    md_to_docx(md, out)


if __name__ == "__main__":
    main()
