#!/usr/bin/env python3
"""Generate English Word tables: Miaolin import supplement request (B2B chatbot)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: list[tuple[str, ...]],
    font_size: float = 10,
) -> None:
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            table.rows[ri].cells[ci].text = val
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(font_size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def build_doc(output_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_heading(doc, "Miaolin — Supplemental data exports for Craveva B2B chatbot", level=1)
    add_para(
        doc,
        "Context: Current sample files under “PROJECT MAOLIN New/” (product, customer, batch inventory, "
        "quotation, period net sales) and the technical target in “zalo chat.txt”. "
        "This document lists what Miaolin should add or confirm in daily/scheduled exports, with rationale.",
    )
    doc.add_paragraph()

    add_heading(doc, "A. Required (highest priority)", level=2)
    add_para(
        doc,
        "Structured sales orders are not replaceable by quotations alone or by period-level sales aggregates.",
        bold=True,
    )
    add_table(
        doc,
        ("Import / structure", "Rationale"),
        (
            (
                "Order export with at least two levels: order header + order line items.",
                "Quotations (“Quote, unit price, inventory”) and period sales (“Last year net sales”) do not answer: "
                "which order, order status, or exactly which SKUs and quantities were on the last real order. "
                "Without order_id and line-level product_code + qty per order, the bot cannot reliably support "
                "CRM-style answers or true “buy again” flows.",
            ),
            (
                "Order header — minimum fields: order_id, customer_code, order_date, status.",
                "status enables answers such as processing / shipped / cancelled; order_date supports “most recent order”.",
            ),
            (
                "Order lines — minimum fields: order_id, product_code, qty, unit_price (or line amount).",
                "Links customer ↔ SKU ↔ quantity — foundation for reorder checks and purchase history.",
            ),
            (
                "Recommended if ERP allows: tracking_number, estimated_delivery.",
                "Reduces “where is my shipment?” tickets; answers can be grounded in data.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "B. Strongly recommended (accuracy & operations)", level=2)
    add_table(
        doc,
        ("Import / field", "Rationale"),
        (
            (
                "Product master: column search_keywords or aliases (synonyms, EN/ZH terms as agreed).",
                "Natural language (“flour”, gluten type, nicknames) maps to SKU more stably than pure LLM guessing. "
                "If Miaolin cannot provide this, Craveva can generate keywords from name + brand + spec, but tuning is slower.",
            ),
            (
                "Inventory rollup by SKU: e.g. product_code, qty_available (optional warehouse_id).",
                "“Craveva full inventory” is batch/warehouse detail — sufficient data but Craveva must SUM. "
                "A pre-aggregated export reduces aggregation risk, lighter jobs, easier audit.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "C. Already covered by current sample exports (no mandatory new file type)", level=2)
    add_table(
        doc,
        ("Data area", "Notes"),
        (
            (
                "Product master (e.g. Craveva product.xlsx)",
                "SKU, name, spec, brand, units, storage — enough for lookup and matching; main gap is optional aliases (see B).",
            ),
            (
                "Customers (e.g. Craveva customer.xlsx)",
                "Customer code and phone support phone-based verification.",
            ),
            (
                "Inventory detail (e.g. Craveva full inventory.xlsx)",
                "Quantities exist; rollup can be done on Craveva side unless Miaolin provides SKU-level totals (see B).",
            ),
            (
                "Quotation / pricing (e.g. Quote, unit price, inventory.xlsx)",
                "Supports pricing context; does not replace structured orders (see A).",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "D. Not expected as ERP spreadsheet columns (handled in Craveva)", level=2)
    add_table(
        doc,
        ("Topic", "Rationale"),
        (
            (
                "LINE User ID / Zalo OA ID mapped to customer_code",
                "Issued by the messaging platform after user interaction — stored and linked in Craveva after OTP/phone verification, "
                "not typically a column on the ERP customer master export.",
            ),
            (
                "auth_username as a separate ERP field",
                "Authentication in practice uses normalized phone + bot session; if phone is already exported, a duplicate "
                "“chat username” column on ERP is optional, not mandatory.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "E. One-line summary for stakeholders", level=2)
    add_para(
        doc,
        "The single most important Miaolin deliverable for this chatbot is a structured order export "
        "(header + lines + status). Existing files mainly support quotation, period sales, and batch stock — "
        "they do not replace real orders. Product aliases (or Craveva-generated) and optional SKU-level stock totals further reduce errors and processing cost.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    out = root / "CHAT" / "Miaolin_Import_Supplement_Request_EN.docx"
    build_doc(out)
    print(f"Wrote: {out}")
