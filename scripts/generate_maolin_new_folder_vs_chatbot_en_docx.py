#!/usr/bin/env python3
"""English Word export for CHAT/maolin_new_folder_vs_chatbot_yeu_cau.md (full tables)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: list[tuple[str, ...]],
    font_size: float = 9,
) -> None:
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            table.rows[ri].cells[ci].text = val
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(font_size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def build_doc(output_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_heading(doc, "Miaolin data (PROJECT MAOLIN New/) vs B2B chatbot requirements", level=1)
    add_para(
        doc,
        "Reference: technical proposal in “zalo chat.txt”. "
        "Source check: first row of each .xlsx (heavy files sampled); “Craveva full inventory.xlsx” uses report title rows — "
        "data column headers appear around row 8 (product code, qty, warehouse name, etc.).",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Data files in PROJECT MAOLIN New/", level=2)
    add_table(
        doc,
        ("File", "Type", "Notes"),
        (
            (
                "Craveva product.xlsx",
                "Product master",
                "Bilingual column headers (e.g. 品號 | SKU, 品名 | Product Name, spec, brand category, shelf life, "
                "inventory mode, storage temperature, expiry date).",
            ),
            (
                "Craveva customer.xlsx",
                "Customers",
                "Customer code, short name, tax ID, sales staff, grade, channel/type/region, delivery address, "
                "TEL_NO (1)(2), payment terms, last transaction, closure date, assigned warehouse name.",
            ),
            (
                "Craveva full inventory.xlsx",
                "Batch / warehouse stock",
                "Batch stock report (批號庫存狀況); data from ~row 8: product code, name, expiry, batch, qty, "
                "remaining shelf days, warehouse name (sample includes last update date).",
            ),
            (
                "Quote, unit price, inventory.xlsx",
                "Quotation + line items",
                "Quote date/no., customer code/name, currency, quote amount, SKU, name, spec, qty, unit price, amount…",
            ),
            (
                "Last year net sales.xlsx",
                "Period net sales",
                "Ship/return date, customer no., product code, net sales qty, net sales amount (local, excl. tax).",
            ),
            (
                "PDF / DOCX",
                "Contracts, BRD",
                "Do not replace structured master data for the bot.",
            ),
        ),
    )
    add_para(
        doc,
        "Not present in this folder: an order export with order_id + status + tracking + order lines as described in “zalo chat.txt”.",
        bold=True,
    )
    doc.add_paragraph()

    add_heading(doc, "2. Quick fit vs chatbot goals", level=2)
    add_table(
        doc,
        ("Area", "In current files", "Notes"),
        (
            (
                "SKU + name + pack + brand",
                "Adequate (Craveva product.xlsx)",
                "Example lookups: product name tokens, 25KG/pack, brand.",
            ),
            (
                "Stock “how much per SKU”",
                "Batch-level detail (Craveva full inventory.xlsx)",
                "Roll up by product code on Craveva side, or request a SKU-level total export.",
            ),
            (
                "Price / quotation",
                "Yes (Quote, unit price, inventory.xlsx)",
                "Quotation data — not equivalent to a full delivery-order lifecycle with status.",
            ),
            (
                "Reorder / “what did I buy” detail",
                "Aggregate only (Last year net sales.xlsx)",
                "Supports “purchased in period”; does not replace per-order + delivery status.",
            ),
            (
                "Aliases / multilingual search keywords",
                "No dedicated column",
                "Craveva can generate from name + brand + spec, or Miaolin adds a column.",
            ),
            (
                "Customer ↔ LINE/Zalo",
                "Not in Excel",
                "Map after phone verification inside Craveva.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "3. Craveva implementation (outside Excel)", level=2)
    add_bullet(
        doc,
        "Sync from PROJECT MAOLIN New/: product, inventory (aggregate if needed), quotation, period sales.",
    )
    add_bullet(
        doc,
        "Product search with confidence; ask clarifying questions when uncertain.",
    )
    add_bullet(doc, "Store aliases / search_keywords in Craveva if Miaolin does not provide a column.")
    add_bullet(
        doc,
        "Conversation memory (recent messages, confirmed SKUs; last order when order feed exists).",
    )
    add_bullet(
        doc,
        "Safety: no cross-customer data leaks; confirmations include SKU + pack; rate-limit phone verification.",
    )
    add_bullet(
        doc,
        "Until a standard order export exists, agree with business whether period sales / confirmed quotes are used as proxies.",
    )
    doc.add_paragraph()

    add_heading(doc, "4. Requested Miaolin exports / columns (summary)", level=2)
    add_bullet(
        doc,
        "Orders: header + order lines (order_id, customer_code, order_date, status, product_code, qty, unit_price; "
        "optional tracking, estimated_delivery).",
    )
    add_bullet(doc, "search_keywords / aliases on product master (or Craveva-generated).")
    add_bullet(doc, "Optional: SKU-level stock totals (sheet/file) to avoid batch SUM on our side.")
    doc.add_paragraph()

    add_heading(doc, "5. Repo helper scripts (header / peek)", level=2)
    p = doc.add_paragraph()
    r = p.add_run(
        "php scripts/read_maolin_new_folder_headers.php\n"
        'php scripts/peek_maolin_sheet.php "PROJECT MAOLIN New/Craveva full inventory.xlsx"'
    )
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "6. Miaolin import supplement — detail (with rationale)", level=2)

    add_heading(doc, "6.1 Required (or equivalent agreed source)", level=3)
    add_table(
        doc,
        ("Import / structure", "Rationale"),
        (
            (
                "Orders — at least two levels: header + line items.",
                "Quotations and period net sales cannot answer which order, order status, or the exact line items of the "
                "last real order. Without order_id and line-level product_code + qty per order, the bot cannot answer "
                "CRM-style or true reorder questions.",
            ),
            (
                "Order header — minimum: order_id, customer_code, order_date, status.",
                "status answers processing / shipped / cancelled; date supports “most recent order”.",
            ),
            (
                "Order lines — minimum: order_id, product_code, qty, unit_price (or line amount).",
                "Links customer ↔ SKU ↔ quantity — basis for reorder and order checks.",
            ),
        ),
    )
    add_para(
        doc,
        "Recommended if ERP allows: tracking_number, estimated_delivery — fewer “where is my shipment?” tickets; grounded answers.",
    )
    doc.add_paragraph()

    add_heading(doc, "6.2 Strongly recommended", level=3)
    add_table(
        doc,
        ("Import / field", "Rationale"),
        (
            (
                "Column search_keywords / aliases on product master.",
                "Customers use natural language (e.g. flour, gluten type, nicknames). Keywords beyond the display name "
                "stabilize SKU mapping vs pure LLM guessing. Miaolin/BI rules preferred; else Craveva generates from name+brand+spec.",
            ),
            (
                "SKU-level stock rollup file/sheet (product_code, qty_available; optional warehouse_id).",
                "Full inventory is batch + warehouse — data is enough but requires SUM. Pre-aggregated export reduces "
                "rollup errors, lighter jobs, easier audit.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "6.3 Already sufficient in current imports (minimal gap)", level=3)
    add_table(
        doc,
        ("Data area", "Short note"),
        (
            (
                "Product master (Craveva product.xlsx)",
                "SKU, name, spec, brand, units, storage — enough for lookup/match; main optional gap is aliases (6.2).",
            ),
            (
                "Customers (Craveva customer.xlsx)",
                "Customer code + phone — enough for phone verification.",
            ),
            (
                "Inventory (Craveva full inventory.xlsx)",
                "Quantities exist in detail form; SUM on Craveva or add rollup file (6.2).",
            ),
            (
                "Quotation (Quote, unit price, inventory.xlsx)",
                "Enough for pricing context; does not replace orders (6.1).",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "6.4 Not expected as ERP columns", level=3)
    add_table(
        doc,
        ("Topic", "Rationale"),
        (
            (
                "LINE User ID / Zalo OA ID ↔ customer_code",
                "Issued by the chat platform after interaction — stored in Craveva, not a typical ERP customer column.",
            ),
            (
                "auth_username as separate ERP export",
                "Auth usually uses normalized phone + bot session; redundant if phone is already exported.",
            ),
        ),
    )
    doc.add_paragraph()

    add_heading(doc, "6.5 Stakeholder one-liner", level=3)
    add_para(
        doc,
        "The top Miaolin add for this B2B chatbot is a structured order export (header + lines + status). "
        "Current files mainly support quotation, period sales, and batch stock — they do not replace real orders. "
        "Product aliases (or Craveva-generated) and optional SKU-level stock totals further reduce errors and cost.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    out = root / "CHAT" / "Miaolin_New_Folder_vs_Chatbot_Requirements_EN.docx"
    build_doc(out)
    print(f"Wrote: {out}")
