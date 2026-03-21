#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_image_block(doc: Document, image_path: Path, title: str, caption: str) -> None:
    add_paragraph(doc, title, bold=True)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.4))
    else:
        add_paragraph(doc, f"[Khong tim thay anh: {image_path.name}]")
    add_paragraph(doc, caption)
    doc.add_paragraph()


def build_report(output_path: Path, report_dir: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_heading(doc, "BAO CAO VAN DE KY THUAT", level=1)
    add_heading(doc, "He thong ERP tich hop AI - Sai lech xac thuc thanh vien (client_code + mobile)", level=2)

    add_heading(doc, "A. Tong quan he thong", level=2)
    add_paragraph(
        doc,
        "He thong hien tai la nen tang ERP tich hop AI Assistant de ho tro nghiep vu theo ngu canh du lieu thuc te "
        "(khach hang, san pham, bao gia, ton kho...).",
    )
    add_paragraph(
        doc,
        "AI duoc ket noi truc tiep voi co so du lieu MySQL thong qua data source da cau hinh, thuc hien truy van va "
        "phan hoi ngay trong luong hoi thoai.",
    )
    add_paragraph(doc, "Dac diem van hanh:")
    add_paragraph(doc, "- AI Agent duoc dung tu quy trinh cau hinh data source, schema analysis, mapping, template va deploy.")
    add_paragraph(doc, "- AI xu ly tac vu nghiep vu bang du lieu truy van truc tiep tu DB.")
    add_paragraph(doc, "- Chuc nang verify thanh vien dua tren cap thong tin client_code va so mobile dang ky.")

    add_heading(doc, "B. Mo ta van de", level=2)
    add_paragraph(
        doc,
        "He thong dang xuat hien loi khong nhat quan du lieu phan hoi cua AI, tap trung ro nhat o chuc nang verify thanh vien.",
    )
    add_paragraph(doc, "- Cung mot luong nghiep vu verify, AI co luc xac thuc thanh cong, co luc khong xac thuc hoac yeu cau xac thuc lap lai.")
    add_paragraph(doc, "- Co truong hop AI thong bao verified du dieu kien doi chieu client_code + mobile khong tra ve ban ghi tuong ung trong DB.")
    add_paragraph(doc, "- Hanh vi verify khong on dinh lam ket qua hoi thoai thay doi bat thuong giua cac lan nhap.")

    add_heading(doc, "C. Phan tich hien tuong", level=2)
    add_paragraph(doc, "1) Nhap sai van pass")
    add_paragraph(doc, "Co hien tuong AI tra trang thai xac thuc thanh cong du dieu kien doi chieu client_code + mobile khong khop du lieu DB.")
    add_paragraph(doc, "2) Nhap dung van fail / khong duoc giu trang thai verify")
    add_paragraph(doc, "Nguoi dung da duoc thong bao xac thuc, nhung o luot sau AI tiep tuc yeu cau nhap lai thong tin verify.")
    add_paragraph(doc, "3) Ket qua khong on dinh theo thoi diem va ngu canh")
    add_paragraph(doc, "- Co luc AI truy van va tra so lieu nghiep vu binh thuong.")
    add_paragraph(doc, "- Co luc cung loai tuong tac lai quay ve buoc verify.")
    add_paragraph(doc, "- Tinh lap lai cua ket qua thap, gay kho kiem chung.")

    add_heading(doc, "D. Bang chung (Evidence)", level=2)
    add_paragraph(doc, "Gia dinh toan bo hinh duoc luu tai thu muc REPORT/.")
    doc.add_paragraph()

    images = [
        ("1.png", "Hinh 1", "Ket noi data source DB thanh cong tren nen tang AI Agent."),
        ("2.png", "Hinh 2", "Schema Analysis hoan tat, cho thay AI da doc cau truc DB."),
        ("3.png", "Hinh 3", "Danh sach entities/relationships duoc nhan dien."),
        ("4.png", "Hinh 4", "Template agent cho nghiep vu verify khach hang da duoc chon."),
        ("5.png", "Hinh 5", "Agent da deploy thanh cong de van hanh thuc te."),
        ("verify_success.png", "Hinh 6", "AI thong bao tai khoan da duoc xac thuc thanh cong sau khi nhap client_code va mobile."),
        ("NO DATA.png", "Hinh 7", "Truy van SQL voi dieu kien client_code + mobile cho ket qua rong (0 ban ghi)."),
        ("keep akking verify.png", "Hinh 8", "AI tiep tuc yeu cau verify, the hien trang thai xac thuc chua on dinh."),
        ("Screenshot 2026-03-21 130745.png", "Hinh 9", "AI tra ket qua so luong san pham 2462 sau trang thai verified."),
        ("Screenshot 2026-03-21 130951.png", "Hinh 10", "Doi chieu SQL truc tiep SELECT COUNT(*) FROM products cung tra 2462."),
        ("6.png", "Hinh 11", "Trong luong hoi gia, AI quay ve yeu cau verify, cho thay hanh vi chua nhat quan."),
    ]

    for file_name, title, caption in images:
        add_image_block(doc, report_dir / file_name, title, caption)

    add_heading(doc, "E. Anh huong", level=2)
    add_paragraph(doc, "- Sai lech xac thuc nguoi dung: co nguy co xac thuc nham doi tuong hoac tu choi sai nguoi dung hop le.")
    add_paragraph(doc, "- Anh huong truc tiep den nghiep vu ban hang/B2B (tra gia, truy cap thong tin theo quyen khach hang).")
    add_paragraph(doc, "- Giam do tin cay cua AI trong moi truong ERP vi ket qua verify khong the du doan theo cung dau vao.")
    add_paragraph(doc, "- Tang rui ro van hanh do ket qua AI khong on dinh giua cac lan tuong tac.")

    add_heading(doc, "F. Ket luan", level=2)
    add_paragraph(
        doc,
        "He thong ERP tich hop AI hien co van de trong yeu ve tinh nhat quan trong xac thuc thanh vien va do on dinh phan hoi du lieu.",
    )
    add_paragraph(
        doc,
        "Bang chung hoi thoai va truy van SQL cho thay ton tai chenh lech giua trang thai verify do AI tra ve va du lieu thuc te trong DB o mot so truong hop.",
    )
    add_paragraph(
        doc,
        "Tong the, van de cot loi la ket qua verify va phan hoi AI chua on dinh, lam suy giam do tin cay nghiep vu cua he thong.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    report_dir = root / "REPORT"
    output_file = report_dir / "ERP_AI_Verify_Issue_Report_VI.docx"
    build_report(output_file, report_dir)
    print(f"Wrote: {output_file}")
