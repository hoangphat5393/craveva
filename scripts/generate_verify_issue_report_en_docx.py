#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_image_block(doc: Document, image_path: Path, title: str, caption: str) -> None:
    add_paragraph(doc, title, bold=True)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.3))
    else:
        add_paragraph(doc, f"[Image not found: {image_path.name}]")
    add_paragraph(doc, caption)
    doc.add_paragraph()


def build_report(output_path: Path, report_dir: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_heading(doc, "Technical Issue Report", level=1)
    add_heading(doc, "ERP + AI Verification Instability (client_code + mobile)", level=2)

    add_heading(doc, "A. System Overview", level=2)
    add_paragraph(
        doc,
        "The current platform is an ERP system integrated with an AI assistant. "
        "The AI is connected directly to the production-like MySQL data source and answers user requests in real time.",
    )
    add_paragraph(
        doc,
        "The implementation flow includes data source connection, schema analysis, entity mapping, prompt template selection, "
        "and deployment as an embedded chat widget.",
    )
    add_bullet(doc, "Business verification flow uses two fields: client_code + registered mobile.")
    add_bullet(doc, "After verification, the AI should allow account-scoped responses (e.g., pricing/account data access).")

    add_heading(doc, "B. Problem Statement", level=2)
    add_paragraph(
        doc,
        "The AI returns inconsistent verification outcomes. The most critical issue is that the assistant can return "
        "'verification success' even when the corresponding record is not found in the database.",
    )
    add_bullet(doc, "Same user journey produces different verification outcomes across attempts.")
    add_bullet(doc, "Verification state is unstable: sometimes retained, sometimes re-requested.")
    add_bullet(doc, "Observed mismatch between AI verification output and SQL ground truth.")

    add_heading(doc, "C. Detailed Verification Anomaly Analysis", level=2)
    add_paragraph(doc, "1) Core contradiction: DB returns no matching user, but AI confirms verification.")
    add_bullet(
        doc,
        "Input pair used in testing: client_code = AVABA, mobile = 0937226422.",
    )
    add_bullet(
        doc,
        "Direct SQL check (JOIN users and client_details) with the same pair returns 0 rows.",
    )
    add_bullet(
        doc,
        "In chat flow, AI still responds with: 'Your account has been successfully verified.'",
    )

    add_paragraph(doc, "2) Behavioral inconsistency after a successful verification message.")
    add_bullet(
        doc,
        "In one sequence, AI acknowledges verified status and answers product count queries (2462).",
    )
    add_bullet(
        doc,
        "In other sequences, AI asks for client_code + mobile again, indicating the verification state is not consistently applied.",
    )

    add_paragraph(doc, "3) Non-deterministic verification behavior.")
    add_bullet(
        doc,
        "Wrong/unsupported inputs may still pass under some interactions.",
    )
    add_bullet(
        doc,
        "Expected-valid flow may fail or reset in later turns.",
    )
    add_bullet(
        doc,
        "The same verification intent does not produce repeatable results.",
    )

    add_heading(doc, "D. Evidence", level=2)
    add_paragraph(doc, "All screenshots are referenced from REPORT/ and embedded below.")

    images = [
        ("1.png", "Figure 1", "AI platform data source list showing successful DB connection."),
        ("2.png", "Figure 2", "Schema analysis completed for TEST DB LOCAL."),
        ("3.png", "Figure 3", "Entities/relationships detected from connected database."),
        ("4.png", "Figure 4", "Template configuration for customer verification conversation flow."),
        ("5.png", "Figure 5", "Agent deployment status is successful."),
        ("verify_success.png", "Figure 6", "Chat shows 'successfully verified' after user submits AVABA + 0937226422."),
        ("NO DATA.png", "Figure 7", "SQL evidence for the same pair returns no matching row (0 result)."),
        ("keep akking verify.png", "Figure 8", "AI asks for verification again in a later interaction."),
        ("Screenshot 2026-03-21 130745.png", "Figure 9", "Post-verify conversation returns product count = 2462."),
        ("Screenshot 2026-03-21 130951.png", "Figure 10", "Direct SQL count from products table also returns 2462."),
        ("6.png", "Figure 11", "In another flow, AI asks for client_code + mobile before providing pricing."),
    ]

    for name, title, caption in images:
        add_image_block(doc, report_dir / name, title, caption)

    add_heading(doc, "E. Impact", level=2)
    add_bullet(doc, "Incorrect user authentication outcome (false positive and false negative risk).")
    add_bullet(doc, "Business workflow risk for B2B pricing and account-specific support.")
    add_bullet(doc, "Trust degradation: users cannot rely on verification consistency.")
    add_bullet(doc, "Operational uncertainty due to mismatch between AI output and database truth.")

    add_heading(doc, "F. Conclusion", level=2)
    add_paragraph(
        doc,
        "The key issue is verification instability in the ERP-integrated AI flow. "
        "Most critically, the system can return verification success even when database evidence shows no matching record.",
    )
    add_paragraph(
        doc,
        "This inconsistency makes the verification result non-reliable for business operations and introduces authentication risk.",
    )
    add_paragraph(
        doc,
        "This report intentionally focuses on issue description and evidence only.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    report_dir = root / "REPORT"
    output_file = report_dir / "ERP_AI_Verification_Issue_Report_EN.docx"
    build_report(output_file, report_dir)
    print(f"Wrote: {output_file}")
